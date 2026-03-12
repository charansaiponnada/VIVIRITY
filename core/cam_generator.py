import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.indian_context import deduplicate_persons, get_cam_financial_rows

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _bold_cell(cell, text: str, font_size: int = 10, color: str = None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


class CAMGenerator:
    VIVRITI_BLUE  = "1B3A6B"
    VIVRITI_GOLD  = "C9A84C"
    LIGHT_BLUE    = "E8F0FA"
    LIGHT_GREY    = "F5F5F5"
    RED_ALERT     = "C0392B"
    GREEN_OK      = "1E8449"
    ORANGE_WARN   = "D68910"

    def __init__(self, output_dir: str = "outputs"):
        os.makedirs(output_dir, exist_ok=True)
        self.output_dir = output_dir
        self.last_pdf_path: str | None = None

    # ================================================================== #
    def generate(
        self,
        company_name:   str,
        financials:     dict,
        research:       dict,
        scoring:        dict,
        cross_ref:      dict,
        manual_notes:   str  = "",
        loan_amount:    str  = "",
        loan_purpose:   str  = "",
    ) -> str:
        print(f"[CAMGenerator] Generating Credit Appraisal Memo for {company_name}...")

        doc = Document()
        self._set_page_margins(doc)
        self._set_default_font(doc)

        self._add_header(doc, company_name, scoring)
        self._add_executive_summary(doc, company_name, scoring, financials, loan_amount, loan_purpose)
        
        # New SWOT Section
        swot = scoring.get("swot", {})
        if swot:
            self._add_swot_analysis(doc, swot)
            
        self._add_company_background(doc, financials, company_name)
        
        # New Pre-Cognitive Section
        from core.risk_engine import detect_precognitive_signals
        precog = detect_precognitive_signals(research, financials)
        if precog:
            self._add_precognitive_signals(doc, precog)
            
        self._add_financial_analysis(doc, financials)
        
        # New Specialized Sections
        self._add_specialized_documents(doc, financials)
        
        self._add_five_cs(doc, scoring)
        self._add_research_intelligence(doc, research, scoring)
        self._add_cross_reference(doc, cross_ref)
        if manual_notes:
            self._add_field_notes(doc, manual_notes)
        self._add_recommendation(doc, scoring)
        self._add_footer(doc)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in company_name if c.isalnum() or c in " _-")[:40]
        filename  = f"CAM_{safe_name}_{timestamp}.docx"
        filepath  = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        self.last_pdf_path = self._try_export_pdf(filepath)
        print(f"[CAMGenerator] CAM generated: {filepath}")
        if self.last_pdf_path:
            print(f"[CAMGenerator] CAM PDF generated: {self.last_pdf_path}")
        return filepath

    # ================================================================== #
    def _add_swot_analysis(self, doc, swot: dict):
        self._section_heading(doc, "SWOT ANALYSIS")
        
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        
        quadrants = [
            ("STRENGTHS", "strengths", "1E8449"),
            ("WEAKNESSES", "weaknesses", "C0392B"),
            ("OPPORTUNITIES", "opportunities", "2471A3"),
            ("THREATS", "threats", "D68910")
        ]
        
        for i, (label, key, color) in enumerate(quadrants):
            row, col = divmod(i, 2)
            cell = table.cell(row, col)
            _set_cell_bg(cell, self.LIGHT_GREY)
            
            p = cell.paragraphs[0]
            run = p.add_run(f" {label}")
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(color)
            
            items = swot.get(key, [])
            for item in items:
                cp = cell.add_paragraph(f"• {item}", style="List Bullet")
                cp.paragraph_format.left_indent = Inches(0.2)
        
        doc.add_paragraph()

    def _add_precognitive_signals(self, doc, precog: list):
        """Add a high-impact table for forward-looking risk signals."""
        self._section_heading(doc, "🚨 PRE-COGNITIVE RISK SIGNALS")
        st = doc.add_paragraph()
        st.add_run("Leading indicators of future credit stress detected through deep intelligence analysis.").italic = True
        
        table = doc.add_table(rows=len(precog)+1, cols=3)
        table.style = "Table Grid"
        
        headers = ["Signal Type", "Detected Signal", "Analytical Insight"]
        for i, h in enumerate(headers):
            hc = table.cell(0, i)
            _set_cell_bg(hc, self.VIVRITI_BLUE)
            _bold_cell(hc, h, color="FFFFFF")
            
        for i, sig in enumerate(precog, 1):
            table.cell(i, 0).text = sig.get("type", "N/A")
            
            # Color coding the signal cell based on impact
            impact = sig.get("impact", "MEDIUM")
            color = self.RED_ALERT if impact == "CRITICAL" else self.ORANGE_WARN if impact == "HIGH" else self.VIVRITI_GOLD
            
            c1 = table.cell(i, 1)
            p = c1.paragraphs[0]
            r = p.add_run(f"[{impact}] {sig.get('signal', 'N/A')}")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(color)
            
            table.cell(i, 2).text = sig.get("insight", "N/A")
            
        doc.add_paragraph()

    def _add_specialized_documents(self, doc, financials: dict):
        """Add sections for ALM, Shareholding, Borrowing, and Portfolio Cuts."""
        
        # 1. Shareholding Pattern
        sh = financials.get("shareholding_pattern", {})
        if sh:
            self._section_heading(doc, "SHAREHOLDING PATTERN")
            table = doc.add_table(rows=4, cols=2)
            table.style = "Table Grid"
            rows = [
                ("Promoter Holding %", f"{sh.get('promoter_holding_percent', 'N/A')}%"),
                ("Public Holding %", f"{sh.get('public_holding_percent', 'N/A')}%"),
                ("Institutional %", f"{sh.get('institutional_holding_percent', 'N/A')}%"),
                ("Pledged Shares %", f"{sh.get('pledged_shares_percent', 'N/A')}%"),
            ]
            for i, (l, v) in enumerate(rows):
                table.cell(i, 0).text = l
                table.cell(i, 1).text = v
            doc.add_paragraph()

        # 2. Borrowing Profile
        bp = financials.get("borrowing_profile", {})
        if bp:
            self._section_heading(doc, "BORROWING PROFILE")
            lenders = bp.get("lender_details", [])
            if lenders:
                table = doc.add_table(rows=len(lenders)+1, cols=4)
                table.style = "Table Grid"
                headers = ["Lender", "Facility", "Limit (Cr)", "O/S (Cr)"]
                for i, h in enumerate(headers):
                    _bold_cell(table.cell(0, i), h)
                for i, l in enumerate(lenders, 1):
                    table.cell(i, 0).text = str(l.get("name", "N/A"))
                    table.cell(i, 1).text = str(l.get("facility", "N/A"))
                    table.cell(i, 2).text = str(l.get("limit", "N/A"))
                    table.cell(i, 3).text = str(l.get("outstanding", "N/A"))
            doc.add_paragraph()

        # 3. ALM (Liquidity GAP)
        alm = financials.get("alm_report", {})
        if alm:
            self._section_heading(doc, "ALM & LIQUIDITY ANALYSIS")
            sl = alm.get("structural_liquidity", {})
            p = doc.add_paragraph()
            p.add_run("Cumulative GAP (Cr): ").bold = True
            p.add_run(str(sl.get("cumulative_gap_crores", "N/A")))
            p.add_run("   |   ")
            p.add_run("Net GAP %: ").bold = True
            p.add_run(f"{sl.get('net_gap_percent', 'N/A')}%")
            doc.add_paragraph()

        # 4. Portfolio Performance
        pc = financials.get("portfolio_cuts", {})
        if pc:
            self._section_heading(doc, "PORTFOLIO PERFORMANCE")
            pq = pc.get("portfolio_quality", {})
            p = doc.add_paragraph()
            p.add_run("GNPA %: ").bold = True
            p.add_run(f"{pq.get('gnpa_percent', 'N/A')}%")
            p.add_run("   |   ")
            p.add_run("Collection Efficiency: ").bold = True
            p.add_run(f"{pq.get('collection_efficiency', 'N/A')}%")
            doc.add_paragraph()

    def _try_export_pdf(self, docx_path: str) -> str | None:
        """Attempt DOCX to PDF conversion; returns PDF path if successful."""
        if not DOCX2PDF_AVAILABLE:
            return None
        try:
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
            docx2pdf_convert(docx_path, pdf_path)
            return pdf_path if os.path.exists(pdf_path) else None
        except Exception as e:
            print(f"[CAMGenerator] PDF conversion skipped: {e}")
            return None

    # ================================================================== #
    def _set_page_margins(self, doc):
        from docx.oxml import OxmlElement
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.0)
            section.right_margin  = Cm(2.0)

    def _set_default_font(self, doc):
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

    # ================================================================== #
    def _add_header(self, doc, company_name: str, scoring: dict):
        rec    = scoring.get("recommendation", {})
        rating = rec.get("rating", scoring.get("risk_score", {}).get("rating", "N/A"))
        decision = rec.get("decision", "N/A")

        # Header table: logo area + title + decision badge
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(2.0)

        left_cell   = table.cell(0, 0)
        mid_cell    = table.cell(0, 1)
        right_cell  = table.cell(0, 2)

        _set_cell_bg(left_cell,  self.VIVRITI_BLUE)
        _set_cell_bg(mid_cell,   self.VIVRITI_BLUE)
        _set_cell_bg(right_cell, self.VIVRITI_BLUE)

        # Left: Company
        lp = left_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = lp.add_run("VIVRITI CAPITAL")
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(self.VIVRITI_GOLD)

        # Mid: Title
        mp = mid_cell.paragraphs[0]
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = mp.add_run("CREDIT APPRAISAL MEMORANDUM\n")
        r1.bold = True; r1.font.size = Pt(14); r1.font.color.rgb = RGBColor(255,255,255)
        r2 = mp.add_run(company_name)
        r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = RGBColor.from_string(self.VIVRITI_GOLD)

        # Right: Decision badge
        decision_color = self.GREEN_OK if decision == "APPROVE" else \
                         self.ORANGE_WARN if decision == "CONDITIONAL_APPROVE" else self.RED_ALERT
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = rp.add_run(f"{decision}\n")
        r3.bold = True; r3.font.size = Pt(13); r3.font.color.rgb = RGBColor.from_string(decision_color)
        r4 = rp.add_run(f"Rating: {rating}")
        r4.bold = True; r4.font.size = Pt(11); r4.font.color.rgb = RGBColor(255,255,255)

        doc.add_paragraph()

    # ================================================================== #
    def _add_executive_summary(self, doc, company_name, scoring, financials, loan_amount, loan_purpose):
        self._section_heading(doc, "EXECUTIVE SUMMARY")

        rec         = scoring.get("recommendation", {})
        risk_score  = scoring.get("risk_score",     {})
        # Use blended score from recommendation (post-ML-blend), fall back to risk_score
        final_score = rec.get("final_score", risk_score.get("final_score", 0))
        rating      = rec.get("rating",  risk_score.get("rating",   "N/A"))
        decision    = rec.get("decision", "N/A")
        amount      = rec.get("recommended_amount_crores", loan_amount or "N/A")
        rate        = rec.get("interest_rate_percent", "N/A")
        tenure      = rec.get("tenure_months", "N/A")

        # Score card table
        table = doc.add_table(rows=2, cols=5)
        table.style = "Table Grid"
        headers = ["Credit Score", "Rating", "Decision", "Recommended Amount", "Interest Rate"]
        values  = [
            f"{final_score}/100",
            str(rating),
            str(decision),
            f"₹{amount} Cr" if amount not in ["N/A", None, ""] else "N/A",
            f"{rate}%" if rate not in ["N/A", None, ""] else "N/A",
        ]

        for i, (h, v) in enumerate(zip(headers, values)):
            hc = table.cell(0, i); vc = table.cell(1, i)
            _set_cell_bg(hc, self.VIVRITI_BLUE)
            _bold_cell(hc, h, font_size=9, color="FFFFFF")
            hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            vc.text = v
            vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            vc.paragraphs[0].runs[0].bold = True
            _set_cell_bg(vc, self.LIGHT_BLUE)

        doc.add_paragraph()

        # Decision rationale
        rationale = rec.get("decision_rationale") or rec.get("rejection_reason") or "See Five Cs analysis below."
        p = doc.add_paragraph()
        p.add_run("Decision Rationale: ").bold = True
        p.add_run(str(rationale))

        # Key conditions
        conditions = rec.get("key_conditions", [])
        if conditions:
            p2 = doc.add_paragraph()
            p2.add_run("Key Conditions: ").bold = True
            p2.add_run("; ".join(str(c) for c in conditions))

        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.add_run(f"Date of Appraisal: ").bold = True
        p3.add_run(datetime.now().strftime("%d %B %Y"))
        if loan_purpose:
            p3.add_run(f"   |   Loan Purpose: ").bold = True
            p3.add_run(str(loan_purpose))

    # ================================================================== #
    def _add_company_background(self, doc, financials: dict, company_name: str):
        self._section_heading(doc, "COMPANY BACKGROUND")
        basic = financials if isinstance(financials, dict) else {}
        entity_type = basic.get("_entity_type", "corporate")

        # Handle directors as list of dicts OR list of strings
        directors_raw = deduplicate_persons(basic.get("directors", []) or [])
        directors_str = ", ".join([
            d.get("name", str(d)) if isinstance(d, dict) else str(d)
            for d in directors_raw
        ]) or "Not extracted"

        promoters_raw = deduplicate_persons(basic.get("promoters", []) or [])
        promoters_str = ", ".join([
            p.get("name", str(p)) if isinstance(p, dict) else str(p)
            for p in promoters_raw
        ]) or "Not extracted"

        rows = [
            ("Company Name",     basic.get("company_name", company_name) or company_name),
            ("CIN",              basic.get("cin",          "Not extracted")),
            ("Entity Type",      str(entity_type).upper()),
            ("Directors",        directors_str),
            ("Promoters",        promoters_str),
            ("Extraction Notes", basic.get("extraction_notes", "—")),
        ]

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(5.0)

        for i, (label, value) in enumerate(rows):
            lc = table.cell(i, 0); vc = table.cell(i, 1)
            _set_cell_bg(lc, self.LIGHT_BLUE)
            lc.text = label; lc.paragraphs[0].runs[0].bold = True
            vc.text = str(value) if value else "—"

        doc.add_paragraph()

    # ================================================================== #
    def _add_financial_analysis(self, doc, financials: dict):
        self._section_heading(doc, "FINANCIAL ANALYSIS")

        f = financials if isinstance(financials, dict) else {}
        entity_type = f.get("_entity_type", "corporate")

        def fmt(val, suffix="Cr"):
            if val is None or val == "": return "N/A"
            try: return f"₹{float(val):,.0f} {suffix}"
            except Exception: return str(val)

        def fmt_pct(val):
            if val is None: return "N/A"
            try: return f"{float(val):.1f}%"
            except Exception: return str(val)

        def fmt_ratio(val):
            if val is None: return "N/A"
            try: return f"{float(val):.2f}x"
            except Exception: return str(val)

        rows = []
        from utils.indian_context import get_cam_financial_rows
        for label, key, fmt_type in get_cam_financial_rows(entity_type):
            val = f.get(key)
            if fmt_type == "cr": display_val = fmt(val)
            elif fmt_type == "pct": display_val = fmt_pct(val)
            elif fmt_type == "ratio": display_val = fmt_ratio(val)
            else: display_val = str(val) if val is not None else "N/A"
            
            # Reasoning logic
            rationale = "Direct extraction"
            if val == 0 and key == "total_borrowings_crores":
                if "debt free" in str(f.get("extraction_notes","")).lower(): rationale = "Inferred: Debt Free"
            elif "computed" in str(f.get("ratios_computed", "")):
                if key in str(f.get("ratios_computed", "")): rationale = "Derived/Computed"
            
            rows.append((label, display_val, rationale))

        table = doc.add_table(rows=len(rows) + 1, cols=3)
        table.style = "Table Grid"

        # Header
        for ci, h in enumerate(["Metric", "Value (FY Latest)", "Rationale / Source"]):
            cell = table.cell(0, ci)
            _set_cell_bg(cell, self.VIVRITI_BLUE)
            _bold_cell(cell, h, color="FFFFFF")

        for i, (label, value, rat) in enumerate(rows, start=1):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = str(value)
            table.cell(i, 2).text = rat

        # Notes section
        if f.get("extraction_notes"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Extraction Notes: ").bold = True
            p.add_run(f["extraction_notes"])

        # Red flags
        red_flags = f.get("red_flags", {})
        if any(red_flags.values()):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("⚠ Red Flags Detected: ").bold = True
            flags_found = [k.replace("_", " ").title() for k, v in red_flags.items() if v]
            p.add_run(", ".join(flags_found))
            p.runs[-1].font.color.rgb = RGBColor.from_string(self.RED_ALERT)

        doc.add_paragraph()

    # ================================================================== #
    def _add_five_cs(self, doc, scoring: dict):
        self._section_heading(doc, "FIVE Cs OF CREDIT ASSESSMENT")

        five_cs    = scoring.get("five_cs",    {})
        risk_score = scoring.get("risk_score", {})

        cs_items = [
            ("Character",  "character",  "25%", "👤"),
            ("Capacity",   "capacity",   "30%", "💰"),
            ("Capital",    "capital",    "20%", "🏛️"),
            ("Collateral", "collateral", "15%", "🔒"),
            ("Conditions", "conditions", "10%", "🌍"),
        ]

        table = doc.add_table(rows=len(cs_items) + 1, cols=4)
        table.style = "Table Grid"

        # Header row
        for i, h in enumerate(["C", "Score", "Weight", "Rationale"]):
            hc = table.cell(0, i)
            _set_cell_bg(hc, self.VIVRITI_BLUE)
            _bold_cell(hc, h, color="FFFFFF")

        for row_idx, (name, key, weight, icon) in enumerate(cs_items, start=1):
            score     = five_cs.get(f"{key}_score",     "N/A")
            rationale = five_cs.get(f"{key}_rationale", "N/A")

            nc = table.cell(row_idx, 0); sc = table.cell(row_idx, 1)
            wc = table.cell(row_idx, 2); rc = table.cell(row_idx, 3)

            if row_idx % 2 == 0:
                for c in [nc, sc, wc, rc]:
                    _set_cell_bg(c, self.LIGHT_BLUE)

            nc.text = f"{icon} {name}"; nc.paragraphs[0].runs[0].bold = True
            sc.text = str(score)
            wc.text = weight
            rc.text = str(rationale)

        doc.add_paragraph()

        # Score breakdown
        p = doc.add_paragraph()
        p.add_run("Weighted Score: ").bold = True
        p.add_run(f"{risk_score.get('weighted_score', 'N/A')}/100")
        p.add_run("   |   ")
        p.add_run("Penalty Applied: ").bold = True
        p.add_run(f"{risk_score.get('penalty_applied', 0)} points")
        p.add_run("   |   ")
        p.add_run("Final Score: ").bold = True
        p.add_run(f"{risk_score.get('final_score', 'N/A')}/100")

        doc.add_paragraph()

    # ================================================================== #
    def _add_research_intelligence(self, doc, research: dict, scoring: dict = None):
        self._section_heading(doc, "SECONDARY RESEARCH INTELLIGENCE")

        sections = [
            ("Company News",      "company_news",       "summary"),
            ("Promoter Background","promoter_background","summary"),
            ("Sector Headwinds",  "sector_headwinds",   "summary"),
            ("Litigation",        "litigation",          "summary"),
            ("Regulatory",        "regulatory",          "summary"),
            ("MCA Signals",       "mca_signals",         "summary"),
        ]

        for label, key, summary_field in sections:
            section_data = research.get(key, {})
            if not section_data:
                continue

            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(section_data.get(summary_field, "No data available.")))
            doc.add_paragraph()

        # Overall sentiment — use computed research_rating from recommendation if available
        overall = research.get("overall_sentiment", {})
        rec_research = (scoring or {}).get("recommendation", {}).get("research_rating", {})
        if overall:
            p = doc.add_paragraph()
            p.add_run("Overall Research Rating: ").bold = True
            # Prefer the structured rubric grade (A/B/C/D) over Gemini's raw sentiment
            rr_grade = rec_research.get("grade") or overall.get("risk_rating", "N/A")
            rr_label = rec_research.get("label") or overall.get("preliminary_recommendation", "N/A")
            p.add_run(f"{rr_grade} — {rr_label}")

            top_risks = overall.get("top_risks", [])
            if top_risks:
                p2 = doc.add_paragraph()
                p2.add_run("Top Risks Identified: ").bold = True
                p2.add_run("; ".join(str(r) for r in top_risks))

        doc.add_paragraph()

    # ================================================================== #
    def _add_cross_reference(self, doc, cross_ref: dict):
        self._section_heading(doc, "CROSS-DOCUMENT FRAUD INTELLIGENCE")

        if not cross_ref.get("cross_reference_performed"):
            doc.add_paragraph(cross_ref.get("reason", "Cross-reference not performed."))
            doc.add_paragraph()
            return

        p = doc.add_paragraph()
        p.add_run("Documents Compared: ").bold = True
        p.add_run(", ".join(cross_ref.get("documents_compared", [])))

        p2 = doc.add_paragraph()
        p2.add_run("Circular Trading Risk: ").bold = True
        p2.add_run(cross_ref.get("circular_trading_risk", "N/A"))

        p3 = doc.add_paragraph()
        p3.add_run("Revenue Inflation Risk: ").bold = True
        p3.add_run(cross_ref.get("revenue_inflation_risk", "N/A"))

        flags = cross_ref.get("flags", [])
        if flags:
            doc.add_paragraph()
            p4 = doc.add_paragraph()
            p4.add_run("⚠ Flags Detected:").bold = True
            for flag in flags:
                fp = doc.add_paragraph(style="List Bullet")
                fp.add_run(f"[{flag.get('severity','?')}] {flag.get('type','?')}: ").bold = True
                fp.add_run(flag.get("description", ""))

        doc.add_paragraph()

    # ================================================================== #
    def _add_field_notes(self, doc, manual_notes: str):
        self._section_heading(doc, "CREDIT OFFICER FIELD OBSERVATIONS")
        p = doc.add_paragraph(manual_notes)
        p.runs[0].italic = True
        doc.add_paragraph()

    # ================================================================== #
    def _add_recommendation(self, doc, scoring: dict):
        self._section_heading(doc, "FINAL RECOMMENDATION")

        rec      = scoring.get("recommendation", {})
        decision = rec.get("decision", "N/A")

        # Decision box
        table = doc.add_table(rows=1, cols=1)
        cell  = table.cell(0, 0)
        color = self.GREEN_OK if decision == "APPROVE" else \
                self.ORANGE_WARN if decision == "CONDITIONAL_APPROVE" else self.RED_ALERT
        _set_cell_bg(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"RECOMMENDATION: {decision}")
        r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(255,255,255)

        doc.add_paragraph()

        rows = [
            ("Decision",           str(decision)),
            ("Credit Rating",      str(rec.get("rating", "N/A"))),
            ("Credit Score",       f"{rec.get('final_score', 'N/A')}/100"),
            ("Recommended Amount", f"₹{rec.get('recommended_amount_crores','N/A')} Cr"),
            ("Interest Rate",      f"{rec.get('interest_rate_percent','N/A')}%"),
            ("Tenure",             f"{rec.get('tenure_months','N/A')} months"),
        ]

        table2 = doc.add_table(rows=len(rows), cols=2)
        table2.style = "Table Grid"
        for i, (label, value) in enumerate(rows):
            lc = table2.cell(i, 0); vc = table2.cell(i, 1)
            _set_cell_bg(lc, self.LIGHT_BLUE)
            lc.text = label; lc.paragraphs[0].runs[0].bold = True
            vc.text = value

        doc.add_paragraph()

        # Rationale
        rationale = rec.get("decision_rationale") or rec.get("rejection_reason") or ""
        if rationale:
            p2 = doc.add_paragraph()
            p2.add_run("Rationale: ").bold = True
            p2.add_run(str(rationale))

        conditions = rec.get("key_conditions", [])
        if conditions:
            doc.add_paragraph()
            doc.add_paragraph().add_run("Conditions Precedent:").bold = True
            for c in conditions:
                cp = doc.add_paragraph(style="List Bullet")
                cp.add_run(str(c))

        doc.add_paragraph()
        disc = doc.add_paragraph(
            "Disclaimer: This Credit Appraisal Memo has been prepared using AI-assisted analysis "
            "by Intelli-Credit (DOMINIX). All recommendations are subject to review and approval "
            "by authorized credit officers as per Vivriti Capital's credit policy."
        )
        disc.runs[0].italic = True
        disc.runs[0].font.size = Pt(8)

    # ================================================================== #
    def _add_footer(self, doc):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            f"Intelli-Credit by DOMINIX  ·  Vivriti Capital Hackathon 2026  ·  "
            f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}"
        )
        r.font.size  = Pt(8)
        r.font.color.rgb = RGBColor.from_string(self.VIVRITI_BLUE)

    # ================================================================== #
    def _section_heading(self, doc, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size  = Pt(12)
        run.font.color.rgb = RGBColor.from_string(self.VIVRITI_BLUE)
        # Underline via border — simpler: just underline
        run.underline = True
        doc.add_paragraph()